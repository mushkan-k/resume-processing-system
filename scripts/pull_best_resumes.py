# Pull the BEST resume for each employee from JobDiva.
# Selection: Most recent EXPERIENCE DATES + substantial content (>500 chars)
# Stubs with only name/email are skipped even if they have a newer upload date.
#
# Usage:
#   python scripts/pull_best_resumes.py --sample 10
#   python scripts/pull_best_resumes.py --all
#   python scripts/pull_best_resumes.py --employee-id 16997161379367

import os
import re
import sys
import time
import base64
import tempfile
import argparse
import logging
import socket
import mysql.connector
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

ROOT_DIR = Path(__file__).resolve().parents[1]
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(ROOT_DIR / "resume_pull.log", encoding="utf-8"),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)
# Force flush after every log line
for h in logging.root.handlers:
    if isinstance(h, logging.StreamHandler):
        h.flush = lambda: sys.stdout.flush()

JOBDIVA_BASE = "https://api.jobdiva.com/apiv2"
MIN_CONTENT_CHARS = 500


def get_db():
    return mysql.connector.connect(
        host=os.getenv("MYSQL_HOST", "127.0.0.1"),
        port=int(os.getenv("MYSQL_PORT", "3305")),
        database=os.getenv("MYSQL_DATABASE", "resume_processing"),
        user=os.getenv("MYSQL_USER", "resume_user"),
        password=os.getenv("MYSQL_PASSWORD", "resume_password")
    )


def authenticate():
    resp = api_get_with_retry(
        f"{JOBDIVA_BASE}/authenticate",
        params={
            "clientid": os.getenv("JOBDIVA_CLIENT_ID"),
            "username": os.getenv("JOBDIVA_USERNAME"),
            "password": os.getenv("JOBDIVA_PASSWORD")
        },
        headers={},
        timeout=30
    )
    if not resp:
        raise Exception("Authentication failed - no response from JobDiva")
    resp.raise_for_status()
    return resp.text.strip().strip('"')


def extract_text_from_bytes(file_bytes, filetype=""):
    """Extract all text including tables from resume file."""
    ext = (filetype or "").lower().strip()
    if not ext:
        if file_bytes[:4] == b'%PDF':
            ext = 'pdf'
        elif file_bytes[:2] == b'PK':
            ext = 'docx'
        else:
            ext = 'docx'

    tmp = tempfile.mktemp(suffix=f".{ext}")
    with open(tmp, 'wb') as f:
        f.write(file_bytes)

    text = ""
    try:
        if ext == "pdf":
            try:
                import pdfplumber
                with pdfplumber.open(tmp) as pdf:
                    pages = []
                    for page in pdf.pages:
                        pt = page.extract_text() or ""
                        for table in page.extract_tables():
                            for row in table:
                                cells = [str(c).strip() for c in row if c]
                                if cells:
                                    pt += "\n" + " | ".join(cells)
                        pages.append(pt)
                    text = "\n".join(pages)
            except ImportError:
                from PyPDF2 import PdfReader
                reader = PdfReader(tmp)
                text = "\n".join(p.extract_text() or "" for p in reader.pages)

        elif ext in ("docx", "doc"):
            import docx
            doc = docx.Document(tmp)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
    except Exception:
        pass

    try:
        os.unlink(tmp)
    except OSError:
        pass
    return text


def find_latest_experience_year(text):
    """Find the most recent year in the resume text (experience indicator)."""
    if not text:
        return 0
    years = re.findall(r'\b(20[0-3]\d)\b', text)
    if re.search(r'\b(present|till date|current|ongoing)\b', text, re.IGNORECASE):
        years.append(str(datetime.now().year))
    return max((int(y) for y in years), default=0)


def score_resume(text, plaintext):
    """
    Score a resume version. Higher = better.
    Returns -1 if the resume is a stub (too short to be useful).
    """
    content = text if len(text) > len(plaintext) else plaintext
    if len(content) < MIN_CONTENT_CHARS:
        return -1  # Stub, skip it

    latest_year = find_latest_experience_year(content)
    content_len = min(len(content), 20000)

    # Most recent experience year dominates, content length breaks ties
    return (latest_year * 10000) + content_len


# Global default socket timeout to prevent SSL hangs on Windows
socket.setdefaulttimeout(20)

# Suppress noisy PDF font warnings
logging.getLogger("pdfplumber").setLevel(logging.ERROR)
logging.getLogger("pdfminer").setLevel(logging.ERROR)

# Thread pool for hard timeout on requests
_executor = ThreadPoolExecutor(max_workers=2)


def _do_get(url, params, headers, timeout):
    """Actual GET request - runs in a thread so we can hard-timeout."""
    session = requests.Session()
    resp = session.get(url, params=params, headers=headers, timeout=timeout)
    session.close()
    return resp


def api_get_with_retry(url, params, headers, timeout=60, retries=3):
    """GET with hard thread-level timeout + retries for SSL hangs on Windows."""
    # Hard timeout = shorter to catch SSL hangs quickly
    hard_timeout = 45
    connect_timeout = (10, min(timeout, 30)) if isinstance(timeout, (int, float)) else timeout

    for attempt in range(retries):
        try:
            future = _executor.submit(_do_get, url, params, headers, connect_timeout)
            resp = future.result(timeout=hard_timeout)
            return resp
        except FuturesTimeout:
            logger.warning(f"  Retry {attempt+1}/{retries}: Hard timeout ({hard_timeout}s) on {url.split('/')[-1]}")
            future.cancel()
            time.sleep(2)
        except (requests.exceptions.ConnectionError,
                requests.exceptions.Timeout,
                requests.exceptions.SSLError,
                ConnectionResetError,
                socket.timeout,
                OSError) as e:
            if attempt < retries - 1:
                wait = 3 * (attempt + 1)
                logger.warning(f"  Retry {attempt+1}/{retries} after {wait}s: {type(e).__name__}")
                time.sleep(wait)
            else:
                logger.error(f"  FAILED after {retries} retries: {type(e).__name__}: {str(e)[:80]}")
                return None
    logger.warning(f"  Giving up on {url.split('/')[-1]} after all retries")
    return None


def find_best_resume(token, candidate_id):
    """Find the best resume: latest experience + real content."""
    headers = {"Authorization": f"Bearer {token}"}
    resp = api_get_with_retry(
        f"{JOBDIVA_BASE}/bi/CandidateResumesDetail",
        params={"candidateId": candidate_id},
        headers=headers, timeout=30
    )
    if not resp or resp.status_code != 200:
        return None
    data = resp.json()
    records = data.get("data", []) if isinstance(data, dict) else (data if isinstance(data, list) else [])
    if not records:
        return None

    best = None
    best_score = -1

    for r in records:
        rid = r.get("RESUMEID")
        if not rid:
            continue

        try:
            detail_resp = api_get_with_retry(
                f"{JOBDIVA_BASE}/bi/ResumeDetail",
                params={"resumeId": rid},
                headers=headers, timeout=60
            )
        except Exception:
            time.sleep(0.5)
            continue

        if not detail_resp or detail_resp.status_code != 200:
            time.sleep(0.3)
            continue

        detail_data = detail_resp.json()
        detail_recs = detail_data.get("data", []) if isinstance(detail_data, dict) else detail_data
        if not detail_recs:
            time.sleep(0.3)
            continue
        rec = detail_recs[0] if isinstance(detail_recs, list) else detail_recs

        plaintext = rec.get("PLAINTEXT", "") or ""
        b64 = rec.get("FILECONTENT_BASE64ENCODED", "") or ""
        filetype = rec.get("FILETYPE", "") or ""

        # Extract text from file
        extracted = ""
        if b64:
            try:
                file_bytes = base64.b64decode(b64)
                extracted = extract_text_from_bytes(file_bytes, filetype)
            except Exception:
                pass

        s = score_resume(extracted, plaintext)
        content = extracted if len(extracted) > len(plaintext) else plaintext

        if s > best_score:
            best_score = s
            best = {
                "resume_id": rid,
                "base64": b64,
                "plaintext": plaintext,
                "filetype": filetype,
                "score": s,
                "latest_year": find_latest_experience_year(content),
                "content_len": len(content),
                "upload_date": r.get("DATECREATED", ""),
            }

        time.sleep(0.3)

    return best


def pull_best_resumes(employee_ids=None, sample_size=None, skip=0):
    """Main: pull best resume for each employee."""
    conn = get_db()
    cur = conn.cursor(dictionary=True)

    if employee_ids:
        placeholders = ",".join(["%s"] * len(employee_ids))
        cur.execute(f"""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name
            FROM resume r JOIN employee e ON r.employee_jobdiva_id = e.id
            WHERE r.employee_jobdiva_id IN ({placeholders})
        """, employee_ids)
    else:
        cur.execute("""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name
            FROM resume r JOIN employee e ON r.employee_jobdiva_id = e.id
            ORDER BY r.employee_jobdiva_id
        """)

    employees = cur.fetchall()
    cur.close()
    conn.close()

    if skip:
        employees = employees[skip:]
        logger.info(f"Skipping first {skip}, starting from #{skip+1}")
    if sample_size:
        employees = employees[:sample_size]

    logger.info(f"Processing {len(employees)} employees...")
    token = authenticate()
    updated = 0
    already_best = 0
    errors = 0

    for i, emp in enumerate(employees):
        cid = emp['employee_jobdiva_id']
        current_rid = emp['resume_id']

        try:
            best = find_best_resume(token, cid)

            if not best or best['score'] < 0:
                already_best += 1
            elif best['resume_id'] != current_rid:
                # Open fresh DB connection for each update (prevents timeout)
                db = get_db()
                db_cur = db.cursor()
                try:
                    db_cur.execute("""
                        UPDATE resume SET resume_id=%s, resume_base64=%s,
                               file_type=%s, summary=%s, updated_at=NOW()
                        WHERE employee_jobdiva_id=%s
                    """, (best['resume_id'], best['base64'],
                          best['filetype'] or 'docx', best['plaintext'], cid))

                    db_cur.execute("""
                        UPDATE employee_skillset SET resume_id=%s WHERE employee_jobdiva_id=%s
                    """, (best['resume_id'], cid))

                    db.commit()
                finally:
                    db_cur.close()
                    db.close()

                updated += 1
                logger.info(
                    f"  [{i+1}/{len(employees)}] UPDATED {emp['employee_name']}: "
                    f"{current_rid} -> {best['resume_id']} "
                    f"(exp year: {best['latest_year']}, {best['content_len']} chars)"
                )
            else:
                already_best += 1

            # Progress every 20 employees + re-auth every 50
            if (i + 1) % 20 == 0:
                logger.info(f"  [{i+1}/{len(employees)}] Progress... {updated} updated, {already_best} best, {errors} err")
            if (i + 1) % 50 == 0:
                token = authenticate()

            time.sleep(0.3)

        except Exception as e:
            errors += 1
            logger.error(f"  [{i+1}/{len(employees)}] ERROR {cid}: {str(e)[:120]}")
            time.sleep(2)
            if errors % 5 == 0:
                try:
                    token = authenticate()
                except Exception:
                    pass
            if errors > 50:
                logger.error("Too many errors, stopping.")
                break

    logger.info(f"\nDONE: {updated} updated, {already_best} already best, {errors} errors")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--sample", type=int)
    parser.add_argument("--all", action="store_true")
    parser.add_argument("--employee-id", type=int)
    parser.add_argument("--skip", type=int, default=0, help="Skip first N employees (resume from where left off)")
    args = parser.parse_args()

    emp_ids = [args.employee_id] if args.employee_id else None
    sample = args.sample if not args.all else None

    if not emp_ids and not args.all and not args.sample:
        parser.print_help()
    else:
        pull_best_resumes(employee_ids=emp_ids, sample_size=sample, skip=args.skip)
