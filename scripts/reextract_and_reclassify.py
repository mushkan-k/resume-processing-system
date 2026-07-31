"""
Re-extract structured data + Reclassify skills for updated resumes.
====================================================================
For employees whose resumes were updated by pull_best_resumes.py:
  1. Reads raw text from resume.summary
  2. Uses Azure OpenAI to extract skills + experience (structured)
  3. Updates resume.skills, resume.experience, resume.education, resume.score
  4. Replaces employee_skillset rows with fresh skills
  5. Classifies each skill as PRIMARY/SECONDARY in the same LLM call

Usage:
    python -m scripts.reextract_and_reclassify --updated     # Only 391 recently-updated
    python -m scripts.reextract_and_reclassify --all         # All employees
    python -m scripts.reextract_and_reclassify --employee-id 12345
    python -m scripts.reextract_and_reclassify --reclassify-only  # Skip extraction, just reclassify
"""
import os
import sys
import json
import time
import argparse
import logging
import re
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime, timedelta
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeout

import mysql.connector
from dotenv import load_dotenv

ROOT_DIR = Path(__file__).resolve().parents[1]
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(ROOT_DIR / "reextract_reclassify.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

DB_CONFIG = {
    "host": os.getenv("MYSQL_HOST", "127.0.0.1"),
    "port": int(os.getenv("MYSQL_PORT", "3305")),
    "database": os.getenv("MYSQL_DATABASE", "resume_processing"),
    "user": os.getenv("MYSQL_USER", "resume_user"),
    "password": os.getenv("MYSQL_PASSWORD", "resume_password"),
    "charset": "utf8mb4",
    "connection_timeout": 30,
}


def get_db():
    return mysql.connector.connect(**DB_CONFIG)


def get_openai_client():
    from openai import AzureOpenAI
    return AzureOpenAI(
        api_key=os.getenv("AZURE_OPENAI_KEY"),
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-08-01-preview")
    )


# =====================================================================
# Prompts
# =====================================================================

EXTRACT_AND_CLASSIFY_PROMPT = """You are an expert HR analyst. Given a resume's raw text, do TWO things:

## TASK 1 — EXTRACT structured data:
- skills: Technical skills, tools, frameworks, platforms, languages, methodologies, techniques
- experience: ALL work history entries with role, company, duration, description
- education: ALL degrees with degree name, institution, year

### SKILL EXTRACTION — What to EXCLUDE (everything else is valid):
EXCLUDE ONLY these:
  - Specific LLM model VERSION NAMES used as a service (GPT-3.5, GPT-4, GPT-4o, Claude 3.5, Gemini 1.5, Mistral-7B, LLaMA-2-70B)
    → These are products you call via API, not skills you possess
    → BUT KEEP the framework/technique if someone fine-tunes or builds with it (e.g. "BERT fine-tuning" = valid skill)
  - Pure soft skills (Communication, Teamwork, Leadership, Time Management)
  
KEEP as valid skills (do NOT exclude):
  - Domain expertise: Machine Learning, Deep Learning, NLP, Computer Vision, Object Detection
  - Techniques/methods: Prompt Engineering, Data Extraction, RAG, Feature Engineering, Regression, Classification
  - Engineering practices: MLOps, CI/CD, Agile, Multi-Agent Systems, Agentic AI
  - Frameworks named after models: LangChain, BERT (when used as a tool), CrewAI, AutoGen
  - Architecture patterns: Multi-modal AI, Feedback Mechanisms, Data Pipelines

## TASK 2 — CLASSIFY each skill as PRIMARY or SECONDARY:
- **PRIMARY**: Core professional identity skill. Used extensively across multiple roles OR central focus of career.
  Typically 3-8 PRIMARY skills per person.
- **SECONDARY**: Supporting/tangential skill. Used briefly, in one short role, or as minor part of duties.

## Classification Rules:
1. 15+ years using a skill → PRIMARY
2. Appears in only 1 short role with no depth → SECONDARY
3. Management/Leadership = PRIMARY only if dedicated manager
4. Soft skills = almost always SECONDARY
5. Consider SENIORITY and RECENCY
6. When in doubt → SECONDARY (PRIMARY = "trust them to interview others on this")

## Resume Text:
{resume_text}

## Output Format (JSON):
{{
  "experience": [
    {{"role": "Software Engineer", "company": "Google", "duration": "Jan 2020 - Present", "description": "Built microservices..."}}
  ],
  "education": [
    {{"degree": "MS Computer Science", "institution": "MIT", "year": "2019"}}
  ],
  "skills": [
    {{"skill": "Python", "type": "PRIMARY", "reason": "Core language used across 10+ years in every role"}},
    {{"skill": "Docker", "type": "SECONDARY", "reason": "Used briefly for containerization in one project"}}
  ]
}}

IMPORTANT:
- Extract ALL skills mentioned (typically 10-30+)
- Do NOT invent skills not in the resume
- Every experience entry must have role, company, duration
- Return ONLY valid JSON, no other text
"""

CLASSIFY_ONLY_PROMPT = """You are an expert HR analyst. Given an employee's work experience and skill list, classify EACH skill as PRIMARY or SECONDARY.

## Definitions:
- **PRIMARY**: Core professional identity skill. Used extensively, demonstrates depth, defines what kind of professional they are. Typically 3-8 per person.
- **SECONDARY**: Supporting skill. Used tangentially, briefly, or as minor part of duties.

## Rules:
1. 15+ years using a skill → PRIMARY
2. Only 1 short role → SECONDARY
3. Management = PRIMARY only if dedicated manager
4. Soft skills = almost always SECONDARY
5. Consider SENIORITY and RECENCY
6. When in doubt → SECONDARY

## Employee Experience:
{experience_text}

## Skills to Classify:
{skills_list}

## Output (JSON array):
[
  {{"skill": "Python", "type": "PRIMARY", "reason": "Core language across all roles"}},
  {{"skill": "Docker", "type": "SECONDARY", "reason": "Used in one role briefly"}}
]

Classify EVERY skill listed. Return ONLY valid JSON array.
"""


def truncate_text(text: str, max_chars: int = 12000) -> str:
    """Truncate resume text to fit context window."""
    if not text or len(text) <= max_chars:
        return text or ""
    # Keep first and last parts
    half = max_chars // 2
    return text[:half] + "\n\n[...TRUNCATED...]\n\n" + text[-half:]


def compute_score(data: Dict) -> float:
    """Compute resume quality score from extracted data."""
    score = 0
    skills = data.get("skills", [])
    experience = data.get("experience", [])
    education = data.get("education", [])
    
    # Skills: up to 30 points
    score += min(len(skills), 30)
    
    # Experience: up to 40 points (5 per role, max 8 roles)
    score += min(len(experience) * 5, 40)
    
    # Education: up to 15 points
    score += min(len(education) * 5, 15)
    
    # Primary skills bonus: up to 15 points
    primary_count = sum(1 for s in skills if isinstance(s, dict) and s.get("type") == "PRIMARY")
    score += min(primary_count * 3, 15)
    
    return min(score, 100)


# =====================================================================
# JobDiva Resume Fetch (fallback for short summaries)
# =====================================================================

def fetch_full_resume_text(employee_id: int) -> Optional[str]:
    """Fetch full resume text from JobDiva API when DB summary is too short."""
    import requests
    import base64
    import tempfile
    
    JOBDIVA_BASE = "https://api.jobdiva.com/apiv2"
    try:
        # Authenticate
        resp = requests.get(f"{JOBDIVA_BASE}/authenticate", params={
            "clientid": os.getenv("JOBDIVA_CLIENT_ID"),
            "username": os.getenv("JOBDIVA_USERNAME"),
            "password": os.getenv("JOBDIVA_PASSWORD")
        }, timeout=30)
        resp.raise_for_status()
        token = resp.text.strip().strip('"')
        
        # Get resume versions
        headers = {"Authorization": f"Bearer {token}"}
        resp = requests.get(f"{JOBDIVA_BASE}/bi/CandidateResumesDetail",
                           params={"candidateId": str(employee_id)},
                           headers=headers, timeout=30)
        if resp.status_code != 200:
            return None
        data = resp.json()
        versions = data.get("data", []) if isinstance(data, dict) else (data if isinstance(data, list) else [])
        if not versions:
            return None
        
        # Get latest resume
        sorted_v = sorted(versions, key=lambda v: v.get("DATECREATED", ""), reverse=True)
        resume_id = sorted_v[0].get("RESUMEID")
        
        resp = requests.get(f"{JOBDIVA_BASE}/bi/ResumeDetail",
                           params={"resumeId": str(resume_id)},
                           headers=headers, timeout=90)
        if resp.status_code != 200:
            return None
        data = resp.json()
        recs = data.get("data", []) if isinstance(data, dict) else data
        if not recs:
            return None
        rec = recs[0] if isinstance(recs, list) else recs
        
        # Try plaintext first
        text = rec.get("PLAINTEXT", "") or ""
        
        # Try extracting from binary if plaintext is short
        b64 = rec.get("FILECONTENT_BASE64ENCODED", "")
        if b64 and len(text) < 500:
            try:
                file_bytes = base64.b64decode(b64)
                filetype = (rec.get("FILETYPE", "") or "").lower().strip()
                if not filetype:
                    if file_bytes[:4] == b'%PDF': filetype = 'pdf'
                    elif file_bytes[:2] == b'PK': filetype = 'docx'
                
                with tempfile.NamedTemporaryFile(suffix=f".{filetype}", delete=False) as f:
                    f.write(file_bytes)
                    tmp_path = f.name
                
                extracted = ""
                if filetype == 'pdf':
                    import pdfplumber
                    with pdfplumber.open(tmp_path) as pdf:
                        for page in pdf.pages:
                            extracted += (page.extract_text() or "") + "\n"
                elif filetype in ('docx', 'doc'):
                    from docx import Document
                    doc = Document(tmp_path)
                    for para in doc.paragraphs:
                        extracted += para.text + "\n"
                    for table in doc.tables:
                        for row in table.rows:
                            extracted += " | ".join(cell.text for cell in row.cells) + "\n"
                os.unlink(tmp_path)
                if len(extracted) > len(text):
                    text = extracted
            except Exception:
                pass
        
        return text.strip() if len(text.strip()) > 200 else None
    except Exception as e:
        logger.warning(f"  JobDiva fetch failed for {employee_id}: {e}")
        return None


def extract_and_classify(client, resume_text: str, model: str) -> Optional[Dict]:
    """Single LLM call: extract + classify."""
    if not resume_text or len(resume_text.strip()) < 50:
        return None
    
    text = truncate_text(resume_text)
    prompt = EXTRACT_AND_CLASSIFY_PROMPT.format(resume_text=text)
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a precise HR data extraction and classification engine. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        raw = response.choices[0].message.content.strip()
        parsed = json.loads(raw)
        return parsed
    except Exception as e:
        logger.error(f"LLM extraction failed: {e}")
        return None


def classify_only(client, experience: List[Dict], skills: List[str], model: str) -> Optional[List[Dict]]:
    """Classify existing skills without re-extraction."""
    if not skills:
        return None
    
    # Format experience
    lines = []
    for i, exp in enumerate(experience or [], 1):
        role = exp.get("role", "Unknown")
        company = exp.get("company", "Unknown")
        duration = exp.get("duration", "Unknown")
        desc = exp.get("description", "")
        lines.append(f"{i}. {role} at {company} ({duration})")
        if desc:
            lines.append(f"   {desc[:200]}")
    
    experience_text = "\n".join(lines) if lines else "No experience data available."
    skills_list = "\n".join(f"- {s}" for s in skills)
    
    prompt = CLASSIFY_ONLY_PROMPT.format(
        experience_text=experience_text,
        skills_list=skills_list
    )
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a precise classification engine. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        raw = response.choices[0].message.content.strip()
        parsed = json.loads(raw)
        
        # Handle wrapped response
        if isinstance(parsed, dict):
            for key, val in parsed.items():
                if isinstance(val, list):
                    parsed = val
                    break
        
        return parsed if isinstance(parsed, list) else None
    except Exception as e:
        logger.error(f"LLM classification failed: {e}")
        return None


def save_extraction(conn, employee_id: int, resume_id: str, data: Dict):
    """Save extracted+classified data to DB."""
    cur = conn.cursor()
    
    # Separate skills list and classification
    skills_raw = data.get("skills", [])
    experience = data.get("experience", [])
    education = data.get("education", [])
    
    # Build clean skill names list for resume.skills column
    skill_names = []
    classifications = []
    for item in skills_raw:
        if isinstance(item, dict):
            name = item.get("skill", "").strip()
            stype = item.get("type", "SECONDARY").upper()
            reason = item.get("reason", "")
            if stype not in ("PRIMARY", "SECONDARY"):
                stype = "SECONDARY"
            if name:
                skill_names.append(name)
                classifications.append({"skill": name, "type": stype, "reason": reason})
        elif isinstance(item, str):
            skill_names.append(item.strip())
            classifications.append({"skill": item.strip(), "type": "SECONDARY", "reason": "No classification provided"})
    
    score = compute_score(data)
    now = datetime.now()
    
    # 1. Update resume table
    cur.execute("""
        UPDATE resume
        SET skills = %s, experience = %s, education = %s, score = %s, updated_at = NOW()
        WHERE employee_jobdiva_id = %s
    """, (
        json.dumps(skill_names),
        json.dumps(experience),
        json.dumps(education),
        score,
        employee_id
    ))
    
    # 2. Delete old skills from employee_skillset
    cur.execute(
        "DELETE FROM employee_skillset WHERE employee_jobdiva_id = %s",
        (employee_id,)
    )
    
    # 3. Insert new skills with classifications
    for cls in classifications:
        try:
            cur.execute("""
                INSERT INTO employee_skillset
                    (resume_id, employee_jobdiva_id, skill, skill_type, classification_reason, classified_at)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (
                resume_id, employee_id, cls["skill"],
                cls["type"], cls["reason"], now
            ))
        except mysql.connector.IntegrityError:
            # Duplicate skill name (case variants) - skip
            pass
    
    conn.commit()
    cur.close()
    return len(classifications)


def save_classification_only(conn, employee_id: int, resume_id: str, classifications: List[Dict]):
    """Update just the classification columns without re-extracting skills."""
    cur = conn.cursor()
    now = datetime.now()
    
    classified_skills = set()
    for item in classifications:
        skill_name = item.get("skill", "").strip()
        skill_type = item.get("type", "SECONDARY").upper()
        reason = item.get("reason", "")
        
        if skill_type not in ("PRIMARY", "SECONDARY"):
            skill_type = "SECONDARY"
        
        if skill_name and skill_name.lower() not in classified_skills:
            cur.execute("""
                UPDATE employee_skillset
                SET skill_type = %s, classification_reason = %s, classified_at = %s
                WHERE employee_jobdiva_id = %s AND LOWER(skill) = LOWER(%s)
            """, (skill_type, reason, now, employee_id, skill_name))
            classified_skills.add(skill_name.lower())
    
    conn.commit()
    cur.close()
    return len(classified_skills)


def get_employees_to_process(conn, mode: str, employee_id: Optional[int] = None, skip: int = 0) -> List[Dict]:
    """Get employees based on mode."""
    cur = conn.cursor(dictionary=True)
    
    if employee_id:
        cur.execute("""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name,
                   r.summary, r.experience, r.skills, r.updated_at
            FROM resume r
            JOIN employee e ON e.id = r.employee_jobdiva_id
            WHERE r.employee_jobdiva_id = %s
        """, (employee_id,))
    elif mode == "updated":
        # Employees updated in last 7 days (from pull_best_resumes)
        cutoff = datetime.now() - timedelta(days=7)
        cur.execute("""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name,
                   r.summary, r.experience, r.skills, r.updated_at
            FROM resume r
            JOIN employee e ON e.id = r.employee_jobdiva_id
            WHERE r.updated_at >= %s
            AND LENGTH(COALESCE(r.summary, '')) > 100
            ORDER BY r.employee_jobdiva_id
        """, (cutoff,))
    elif mode == "unclassified":
        cur.execute("""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name,
                   r.summary, r.experience, r.skills, r.updated_at
            FROM resume r
            JOIN employee e ON e.id = r.employee_jobdiva_id
            WHERE LENGTH(COALESCE(r.summary, '')) > 100
            AND r.employee_jobdiva_id NOT IN (
                SELECT DISTINCT employee_jobdiva_id FROM employee_skillset
            )
            ORDER BY r.employee_jobdiva_id
        """)
    else:  # all
        cur.execute("""
            SELECT r.employee_jobdiva_id, r.resume_id, e.employee_name,
                   r.summary, r.experience, r.skills, r.updated_at
            FROM resume r
            JOIN employee e ON e.id = r.employee_jobdiva_id
            WHERE LENGTH(COALESCE(r.summary, '')) > 100
            ORDER BY r.employee_jobdiva_id
        """)
    
    employees = cur.fetchall()
    cur.close()
    
    if skip > 0:
        employees = employees[skip:]
    
    return employees


def process_employees(
    mode: str = "updated",
    reclassify_only: bool = False,
    employee_id: Optional[int] = None,
    skip: int = 0,
    batch_pause: float = 1.0
):
    """Main processing loop."""
    conn = get_db()
    client = get_openai_client()
    model = os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4o-mini")
    
    employees = get_employees_to_process(conn, mode, employee_id, skip)
    conn.close()
    
    total = len(employees)
    logger.info(f"Processing {total} employees (mode={mode}, reclassify_only={reclassify_only})")
    
    if total == 0:
        logger.info("Nothing to process!")
        return
    
    success = 0
    errors = 0
    primary_total = 0
    secondary_total = 0
    
    for i, emp in enumerate(employees, 1):
        emp_id = emp["employee_jobdiva_id"]
        resume_id = emp["resume_id"]
        name = emp["employee_name"]
        summary = emp.get("summary") or ""
        
        try:
            if reclassify_only:
                # Just reclassify existing skills
                conn2 = get_db()
                cur2 = conn2.cursor(dictionary=True)
                cur2.execute(
                    "SELECT skill FROM employee_skillset WHERE employee_jobdiva_id = %s",
                    (emp_id,)
                )
                skills = [r["skill"] for r in cur2.fetchall()]
                cur2.close()
                
                experience = json.loads(emp["experience"]) if emp.get("experience") else []
                
                if not skills:
                    conn2.close()
                    logger.warning(f"  [{i}/{total}] {name}: No skills, skipping")
                    continue
                
                # If no experience data, use summary text as context
                if not experience and summary:
                    # Fake one experience entry from summary
                    experience = [{"role": "Professional", "company": "Various", "duration": "Career", "description": summary[:2000]}]
                
                result = classify_only(client, experience, skills, model)
                
                if result:
                    count = save_classification_only(conn2, emp_id, resume_id, result)
                    p = sum(1 for r in result if r.get("type", "").upper() == "PRIMARY")
                    s = count - p
                    primary_total += p
                    secondary_total += s
                    logger.info(f"  [{i}/{total}] {name}: {p} PRIMARY, {s} SECONDARY")
                    success += 1
                else:
                    logger.warning(f"  [{i}/{total}] {name}: Classification returned None")
                    errors += 1
                
                conn2.close()
            else:
                # Full extraction + classification
                text_for_llm = summary
                if len(summary.strip()) < 500:
                    logger.info(f"  [{i}/{total}] {name}: Summary short ({len(summary)} chars), fetching full resume from JobDiva...")
                    full_text = fetch_full_resume_text(emp_id)
                    if full_text:
                        text_for_llm = full_text
                        logger.info(f"  [{i}/{total}] {name}: Got {len(full_text)} chars from JobDiva")
                    elif len(summary.strip()) < 100:
                        logger.warning(f"  [{i}/{total}] {name}: No text available, skipping")
                        continue
                
                result = extract_and_classify(client, text_for_llm, model)
                
                if result and result.get("skills"):
                    conn2 = get_db()
                    count = save_extraction(conn2, emp_id, resume_id, result)
                    conn2.close()
                    
                    skills_data = result.get("skills", [])
                    p = sum(1 for s in skills_data if isinstance(s, dict) and s.get("type", "").upper() == "PRIMARY")
                    s = count - p
                    primary_total += p
                    secondary_total += s
                    logger.info(f"  [{i}/{total}] {name}: Extracted {count} skills ({p} PRIMARY, {s} SECONDARY)")
                    success += 1
                else:
                    logger.warning(f"  [{i}/{total}] {name}: Extraction returned no skills")
                    errors += 1
            
            # Rate limit
            time.sleep(batch_pause)
            
            # Progress log every 50
            if i % 50 == 0:
                logger.info(f"  Progress: {i}/{total} | {success} ok, {errors} err | {primary_total} PRIMARY, {secondary_total} SECONDARY")
                
        except Exception as e:
            logger.error(f"  [{i}/{total}] {name} (ID={emp_id}): FAILED - {e}")
            errors += 1
            time.sleep(2)  # Extra pause on error
    
    logger.info(f"\nDONE: {success} success, {errors} errors")
    logger.info(f"  Total skills: {primary_total} PRIMARY + {secondary_total} SECONDARY = {primary_total + secondary_total}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Re-extract and reclassify employee skills")
    parser.add_argument("--updated", action="store_true", help="Process only recently-updated resumes")
    parser.add_argument("--all", action="store_true", help="Process ALL employees")
    parser.add_argument("--unclassified", action="store_true", help="Process only employees with NO skills in employee_skillset")
    parser.add_argument("--reclassify-only", action="store_true", help="Skip extraction, just reclassify existing skills")
    parser.add_argument("--employee-id", type=int, help="Process a specific employee")
    parser.add_argument("--skip", type=int, default=0, help="Skip first N employees")
    parser.add_argument("--pause", type=float, default=1.0, help="Seconds between API calls")
    
    args = parser.parse_args()
    
    if not any([args.updated, args.all, args.employee_id, args.unclassified]):
        parser.error("Must specify --updated, --all, --unclassified, or --employee-id")
    
    mode = "unclassified" if args.unclassified else ("all" if args.all else "updated")
    
    process_employees(
        mode=mode,
        reclassify_only=args.reclassify_only,
        employee_id=args.employee_id,
        skip=args.skip,
        batch_pause=args.pause
    )
