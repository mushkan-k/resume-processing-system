# Verify resume quality AFTER pull_best_resumes.py finishes.
# Cross-checks old vs new resumes to ensure no regressions (stubs, empty, lost content).
#
# What we check:
#   1. Every employee has substantial content (>500 chars)
#   2. Every employee has identifiable experience dates
#   3. Compare old vs new from the pull log to catch downgrades
#   4. Flag any suspicious cases for manual review
#
# Usage:
#   python scripts/verify_resume_quality.py
#   python scripts/verify_resume_quality.py --fix   (revert bad updates)

import os
import re
import sys
import base64
import tempfile
import argparse
import logging
import mysql.connector
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

ROOT_DIR = Path(__file__).resolve().parents[1]
load_dotenv(ROOT_DIR / ".env")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.StreamHandler(sys.stdout)]
)
logger = logging.getLogger(__name__)

# Suppress noisy PDF font warnings
logging.getLogger("pdfplumber").setLevel(logging.ERROR)
logging.getLogger("pdfminer").setLevel(logging.ERROR)
import warnings
warnings.filterwarnings("ignore", message=".*FontBBox.*")


def get_db():
    return mysql.connector.connect(
        host=os.getenv("MYSQL_HOST", "127.0.0.1"),
        port=int(os.getenv("MYSQL_PORT", "3305")),
        database=os.getenv("MYSQL_DATABASE", "resume_processing"),
        user=os.getenv("MYSQL_USER", "resume_user"),
        password=os.getenv("MYSQL_PASSWORD", "resume_password")
    )


def extract_text_from_base64(b64_content, filetype=""):
    """Extract all text from a base64-encoded resume file (same logic as pull script)."""
    if not b64_content:
        return ""

    try:
        file_bytes = base64.b64decode(b64_content)
    except Exception:
        return ""

    ext = (filetype or "").lower().strip()
    if not ext:
        if file_bytes[:4] == b'%PDF':
            ext = 'pdf'
        elif file_bytes[:2] == b'PK':
            ext = 'docx'
        else:
            ext = 'docx'

    tmp = tempfile.mktemp(suffix=f".{ext}")
    with open(tmp, 'wb') as f:
        f.write(file_bytes)

    text = ""
    try:
        if ext == "pdf":
            try:
                import pdfplumber
                import warnings
                warnings.filterwarnings("ignore")
                with pdfplumber.open(tmp) as pdf:
                    pages = []
                    for page in pdf.pages:
                        pt = page.extract_text() or ""
                        for table in page.extract_tables():
                            for row in table:
                                cells = [str(c).strip() for c in row if c]
                                if cells:
                                    pt += "\n" + " | ".join(cells)
                        pages.append(pt)
                    text = "\n".join(pages)
            except ImportError:
                pass

        elif ext in ("docx", "doc"):
            import docx
            doc = docx.Document(tmp)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts)
    except Exception:
        pass

    try:
        os.unlink(tmp)
    except OSError:
        pass
    return text


def find_experience_years(text):
    """Find all years mentioned in context that looks like experience."""
    if not text:
        return []
    years = [int(y) for y in re.findall(r'\b(20[0-3]\d)\b', text)]
    if re.search(r'\b(present|till date|current|ongoing)\b', text, re.IGNORECASE):
        years.append(datetime.now().year)
    return sorted(set(years))


def has_skills_content(text):
    """Check if the text has skill-like keywords (technologies, tools, etc)."""
    skill_indicators = [
        r'\b(python|java|javascript|sql|aws|azure|docker|kubernetes)\b',
        r'\b(react|angular|node|spring|django|flask)\b',
        r'\b(machine learning|data science|devops|cloud|agile)\b',
        r'\b(management|leadership|project manager|architect)\b',
        r'\b(excel|powerpoint|tableau|power bi|salesforce)\b',
        r'\b(certified|certification|PMP|AWS|GCP)\b',
    ]
    count = 0
    for pattern in skill_indicators:
        if re.search(pattern, text, re.IGNORECASE):
            count += 1
    return count


def parse_pull_log():
    """Parse the pull log to get old->new resume mappings."""
    log_path = ROOT_DIR / "resume_pull_stdout.log"
    if not log_path.exists():
        log_path = ROOT_DIR / "resume_pull.log"
    if not log_path.exists():
        return {}

    updates = {}
    pattern = r'UPDATED (.+?): (\S+) -> (\S+) \(exp year: (\d+), (\d+) chars\)'
    with open(log_path, 'r', encoding='utf-8') as f:
        for line in f:
            m = re.search(pattern, line)
            if m:
                name = m.group(1)
                old_rid = m.group(2)
                new_rid = m.group(3)
                exp_year = int(m.group(4))
                content_len = int(m.group(5))
                updates[name] = {
                    'old_resume_id': old_rid,
                    'new_resume_id': new_rid,
                    'exp_year': exp_year,
                    'content_len': content_len
                }
    return updates


def verify_all_resumes(fix_mode=False):
    """Main verification: check every employee's current resume quality."""
    conn = get_db()
    cur = conn.cursor(dictionary=True)

    # Get all employees with their current resume
    cur.execute("""
        SELECT e.id AS employee_id, e.employee_name,
               r.resume_id, r.resume_base64, r.file_type, r.summary
        FROM employee e
        JOIN resume r ON e.id = r.employee_jobdiva_id
        ORDER BY e.employee_name
    """)
    employees = cur.fetchall()
    logger.info(f"Verifying {len(employees)} employee resumes...")

    # Parse the pull log for old->new comparison
    pull_updates = parse_pull_log()
    logger.info(f"Found {len(pull_updates)} updates in pull log for comparison")

    # Quality checks
    issues = {
        'stubs': [],         # < 500 chars (name/email only)
        'no_experience': [], # No year dates found
        'no_skills': [],     # No skill keywords found
        'downgraded': [],    # New resume worse than old (from log comparison)
    }
    good = 0
    total_checked = 0

    for emp in employees:
        total_checked += 1
        name = emp['employee_name']
        b64 = emp['resume_base64'] or ""
        filetype = emp['file_type'] or ""
        plaintext = emp['summary'] or ""

        # Extract full text from stored resume
        extracted = extract_text_from_base64(b64, filetype) if b64 else ""
        content = extracted if len(extracted) > len(plaintext) else plaintext
        content_len = len(content)

        # CHECK 1: Stub detection
        if content_len < 500:
            issues['stubs'].append({
                'name': name,
                'employee_id': emp['employee_id'],
                'resume_id': emp['resume_id'],
                'content_len': content_len,
                'preview': content[:100].replace('\n', ' ')
            })
            continue

        # CHECK 2: Experience dates
        years = find_experience_years(content)
        if not years:
            issues['no_experience'].append({
                'name': name,
                'employee_id': emp['employee_id'],
                'resume_id': emp['resume_id'],
                'content_len': content_len
            })

        # CHECK 3: Skills content
        skill_score = has_skills_content(content)
        if skill_score == 0:
            issues['no_skills'].append({
                'name': name,
                'employee_id': emp['employee_id'],
                'resume_id': emp['resume_id'],
                'content_len': content_len
            })

        # CHECK 4: Downgrade detection (compare with pull log)
        if name in pull_updates:
            update = pull_updates[name]
            if update['content_len'] < 500 or update['exp_year'] < 2020:
                issues['downgraded'].append({
                    'name': name,
                    'employee_id': emp['employee_id'],
                    'old_resume_id': update['old_resume_id'],
                    'new_resume_id': update['new_resume_id'],
                    'exp_year': update['exp_year'],
                    'content_len': update['content_len']
                })

        if content_len >= 500 and years and skill_score > 0:
            good += 1

        if total_checked % 200 == 0:
            logger.info(f"  Checked {total_checked}/{len(employees)}...")

    # REPORT
    print("\n" + "=" * 70)
    print("RESUME QUALITY VERIFICATION REPORT")
    print("=" * 70)
    print(f"\nTotal employees checked: {total_checked}")
    print(f"Good quality resumes:   {good} ({good*100//total_checked}%)")
    print(f"\n--- ISSUES FOUND ---")

    print(f"\n🚨 STUBS (< 500 chars, likely name/email only): {len(issues['stubs'])}")
    for item in issues['stubs'][:20]:
        print(f"   {item['name']} | {item['content_len']} chars | ID: {item['resume_id']}")
        print(f"      Preview: {item['preview']}")

    print(f"\n⚠️  NO EXPERIENCE DATES found: {len(issues['no_experience'])}")
    for item in issues['no_experience'][:10]:
        print(f"   {item['name']} | {item['content_len']} chars | ID: {item['resume_id']}")

    print(f"\n⚠️  NO SKILL KEYWORDS found: {len(issues['no_skills'])}")
    for item in issues['no_skills'][:10]:
        print(f"   {item['name']} | {item['content_len']} chars | ID: {item['resume_id']}")

    print(f"\n🔄 POTENTIAL DOWNGRADES (from pull log): {len(issues['downgraded'])}")
    for item in issues['downgraded'][:10]:
        print(f"   {item['name']} | {item['old_resume_id']} -> {item['new_resume_id']}")
        print(f"      Exp year: {item['exp_year']}, Content: {item['content_len']} chars")

    # Summary stats
    print(f"\n{'=' * 70}")
    print(f"SUMMARY:")
    print(f"  ✅ Good:          {good}")
    print(f"  🚨 Stubs:         {len(issues['stubs'])}")
    print(f"  ⚠️  No dates:      {len(issues['no_experience'])}")
    print(f"  ⚠️  No skills:     {len(issues['no_skills'])}")
    print(f"  🔄 Downgrades:    {len(issues['downgraded'])}")
    print(f"{'=' * 70}")

    # Fix mode: revert bad updates
    if fix_mode and issues['stubs']:
        print(f"\n🔧 FIX MODE: Attempting to revert {len(issues['stubs'])} stub resumes...")
        reverted = 0
        for item in issues['stubs']:
            name = item['name']
            if name in pull_updates:
                old_rid = pull_updates[name]['old_resume_id']
                new_rid = pull_updates[name]['new_resume_id']
                # Revert to old resume ID
                cur.execute("""
                    UPDATE resume SET resume_id=%s, updated_at=NOW()
                    WHERE employee_jobdiva_id=%s AND resume_id=%s
                """, (old_rid, item['employee_id'], new_rid))
                cur.execute("""
                    UPDATE employee_skillset SET resume_id=%s
                    WHERE employee_jobdiva_id=%s AND resume_id=%s
                """, (old_rid, item['employee_id'], new_rid))
                if cur.rowcount > 0:
                    reverted += 1
                    logger.info(f"  REVERTED {name}: {new_rid} -> {old_rid}")
        conn.commit()
        print(f"  ✅ Reverted {reverted} employees to their previous resume")

    cur.close()
    conn.close()
    return issues


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Verify resume quality after pull")
    parser.add_argument("--fix", action="store_true", help="Revert bad updates (stubs)")
    args = parser.parse_args()
    verify_all_resumes(fix_mode=args.fix)
